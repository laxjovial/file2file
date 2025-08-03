import streamlit as st
import pandas as pd
from io import BytesIO
import os
from docx import Document
from pdf2docx import Converter # Make sure pdf2docx is installed: pip install pdf2docx
import pdfplumber # Make sure pdfplumber is installed: pip install pdfplumber
from reportlab.pdfgen import canvas # Make sure reportlab is installed: pip install reportlab
from reportlab.lib.pagesizes import letter
import pypandoc # Make sure pypandoc is installed: pip install pypandoc

# Set Streamlit page configuration
st.set_page_config(page_title="File2File Converter", layout="centered")
st.title("📁 File2File Converter")
st.markdown("Convert between PDF, DOCX, TXT, CSV, XLS, XLSX. Batch uploads supported. Real PDF export.")

# Define supported file formats
doc_types = ["pdf", "docx", "txt"]
sheet_types = ["csv", "xls", "xlsx"]
all_types = doc_types + sheet_types

# User selects source and target formats
source_format = st.selectbox("From format", all_types)
target_format = st.selectbox("To format", [f for f in all_types if f != source_format])

# User uploads files
uploaded_files = st.file_uploader(
    "Upload files",
    type=[source_format], # Restrict upload type to selected source format
    accept_multiple_files=True
)

# Optional output file name input
custom_name = st.text_input("Optional: base name for output file(s)", "converted")

# Function to preview uploaded files
def preview_file(file, file_type):
    st.subheader("🔍 Preview")
    # Reset file pointer to the beginning before reading for preview
    file.seek(0) 

    if file_type == "txt":
        text = file.read().decode("utf-8")
        st.text(text[:1000] + ("..." if len(text) > 1000 else ""))
    elif file_type == "csv":
        df = pd.read_csv(file)
        st.dataframe(df.head())
    elif file_type in ["xls", "xlsx"]:
        df = pd.read_excel(file)
        st.dataframe(df.head())
    elif file_type == "docx":
        # For docx, we need to save to a BytesIO object first for python-docx to read
        doc_bytes = BytesIO(file.read())
        doc = Document(doc_bytes)
        text = "\n".join([p.text for p in doc.paragraphs])
        st.text(text[:1000] + ("..." if len(text) > 1000 else ""))
    elif file_type == "pdf":
        # For pdf, we need to save to a BytesIO object first for pdfplumber to read
        pdf_bytes = BytesIO(file.read())
        with pdfplumber.open(pdf_bytes) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            st.text(text[:1000] + ("..." if len(text) > 1000 else ""))

# Function for document (PDF, DOCX, TXT) conversions
def convert_doc_file(uploaded_file, source, target):
    result = BytesIO()
    # Reset file pointer to the beginning before any conversion operation
    # This is crucial because preview_file might have consumed the stream
    uploaded_file.seek(0) 

    # Generate unique temporary filenames to avoid conflicts
    unique_id = os.urandom(8).hex()
    temp_input_path = f"temp_input_{unique_id}.{source}"
    temp_output_path = f"temp_output_{unique_id}.{target}"

    try:
        if source == "pdf":
            # Save the uploaded PDF to a temporary file for pdf2docx
            with open(temp_input_path, "wb") as f:
                f.write(uploaded_file.read())

            if target == "docx":
                cv = Converter(temp_input_path)
                cv.convert(temp_output_path, start=0, end=None)
                cv.close()
                with open(temp_output_path, "rb") as f:
                    result.write(f.read())
            elif target == "txt":
                # pdfplumber can directly open BytesIO, no need for temp file
                with pdfplumber.open(uploaded_file) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                    result.write(text.encode("utf-8"))

        elif source == "docx":
            # Save the uploaded DOCX to a temporary file for pypandoc
            with open(temp_input_path, "wb") as f:
                f.write(uploaded_file.read())

            if target == "pdf":
                # pypandoc requires input and output file paths
                pypandoc.convert_file(temp_input_path, "pdf", outputfile=temp_output_path)
                with open(temp_output_path, "rb") as f:
                    result.write(f.read())
            elif target == "txt":
                # python-docx can directly open BytesIO
                doc_bytes = BytesIO(uploaded_file.read())
                doc = Document(doc_bytes)
                text = "\n".join([p.text for p in doc.paragraphs])
                result.write(text.encode("utf-8"))

        elif source == "txt":
            text = uploaded_file.read().decode("utf-8")
            if target == "pdf":
                c = canvas.Canvas(result, pagesize=letter)
                y = 750 # Starting Y position for text
                for line in text.splitlines():
                    # Draw string, move down for next line
                    c.drawString(50, y, line)
                    y -= 15
                    # Add new page if current page is full
                    if y < 50:
                        c.showPage()
                        y = 750
                c.save() # Save the PDF content to the BytesIO object
            elif target == "docx":
                doc = Document()
                for line in text.splitlines():
                    doc.add_paragraph(line)
                doc.save(result) # Save the DOCX content to the BytesIO object

    finally:
        # Clean up temporary files regardless of conversion success or failure
        if os.path.exists(temp_input_path):
            os.remove(temp_input_path)
        if os.path.exists(temp_output_path):
            os.remove(temp_output_path)

    # Reset the BytesIO object's pointer before returning it for download
    result.seek(0)
    return result

# Function for spreadsheet (CSV, XLS, XLSX) conversions
def convert_sheet_file(uploaded_file, source, target):
    result = BytesIO()
    # Reset file pointer to the beginning before any conversion operation
    uploaded_file.seek(0)

    if source == "csv":
        df = pd.read_csv(uploaded_file)
    else: # source is xls or xlsx
        df = pd.read_excel(uploaded_file)

    if target == "csv":
        df.to_csv(result, index=False)
    else: # target is xls or xlsx
        # For Excel, use 'openpyxl' engine for xlsx, 'xlrd' for xls (though xlrd is deprecated for xls)
        # It's generally safer to just use openpyxl for both if possible, or handle xls specifically
        df.to_excel(result, index=False, engine="openpyxl") 

    # Reset the BytesIO object's pointer before returning it for download
    result.seek(0)
    return result

# Main conversion and download section
if uploaded_files:
    for idx, uploaded_file in enumerate(uploaded_files):
        st.divider()
        st.subheader(f"📄 File {idx + 1}: {uploaded_file.name}")
        
        # Preview the file content
        preview_file(uploaded_file, source_format)

        with st.spinner("Converting..."):
            output = None # Initialize output to None
            if source_format in doc_types and target_format in doc_types:
                output = convert_doc_file(uploaded_file, source_format, target_format)
            elif source_format in sheet_types and target_format in sheet_types:
                output = convert_sheet_file(uploaded_file, source_format, target_format)
            else:
                st.error("❌ Cross-type conversions (e.g., DOCX → CSV) not supported.")
                continue # Skip to the next file if conversion is not supported

        # If output was generated successfully
        if output:
            # Determine the base name for the output file
            file_base = custom_name if custom_name else os.path.splitext(uploaded_file.name)[0]
            
            # Construct the download filename
            if len(uploaded_files) > 1:
                download_name = f"{file_base}_{idx + 1}.{target_format}"
            else:
                download_name = f"{file_base}.{target_format}"

            st.success(f"✅ Done: {download_name}")
            # Provide a download button for the converted file
            st.download_button("⬇️ Download", data=output, file_name=download_name)
